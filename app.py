import os
import uuid
import json
import torch
from pathlib import Path
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO

# Import our robust pipeline
from src.pipeline import FullPagePashtoRecognition
from finetune_user import train_on_user_data
import cv2
import threading

app = Flask(__name__)
app.secret_key = "pashto_ocr_secret"

# Storage
UPLOAD_DIR = Path("static/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///ocr_v2.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

class ScanHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200))
    recognized_text = db.Column(db.Text)
    confidence = db.Column(db.Integer)
    timestamp = db.Column(db.DateTime, default=db.func.current_timestamp())

# Global Pipeline Instance
pipeline = None

def initialize_storage_and_db():
    global pipeline
    with app.app_context():
        db.create_all()
    if pipeline is None:
        pipeline = FullPagePashtoRecognition()
    print("Loaded PashtoCRNN Success.")
    print("Loaded YOLOv8 Segmenter Success.")

def run_ocr_text(img_path):
    global pipeline
    if pipeline is None:
        pipeline = FullPagePashtoRecognition()
    results, confidence = pipeline.process_page(str(img_path))
    return "\n".join(results), confidence

def scan_to_view(scan):
    text = scan.recognized_text or ""
    # Calculate real file size from the uploads directory
    saved_path = UPLOAD_DIR / scan.filename
    f_size = os.path.getsize(saved_path) if saved_path.exists() else 0
    
    return {
        "id": scan.id,
        "filename": scan.filename,
        "text": text,
        "confidence": scan.confidence,
        "word_count": len(text.split()),
        "char_count": len(text),
        "file_size": f_size,
        "date": scan.timestamp.strftime("%Y-%m-%d %H:%M")
    }

@app.route("/")
def dashboard():
    total_scans = ScanHistory.query.count()
    avg_conf = db.session.query(db.func.avg(ScanHistory.confidence)).scalar() or 0
    latest = ScanHistory.query.order_by(ScanHistory.timestamp.desc()).first()
    recent_scans = ScanHistory.query.order_by(ScanHistory.timestamp.desc()).limit(5).all()
    recent_scans = [scan_to_view(s) for s in recent_scans]
    
    return render_template(
        "index.html",
        total_scans=total_scans,
        avg_confidence=int(avg_conf),
        latest=latest,
        recent_scans=recent_scans,
    )

@app.route("/new-scan")
def new_scan():
    return render_template("new_scan.html")

@app.route("/predict", methods=["POST"])
def predict():
    uploaded = request.files.get("document")
    if not uploaded or not uploaded.filename:
        flash("Please select an image file.", "error")
        return redirect(url_for("new_scan"))

    original_name = secure_filename(uploaded.filename)
    scan_id = uuid.uuid4().hex[:12]
    saved_name = f"{scan_id}_{original_name}"
    saved_path = UPLOAD_DIR / saved_name
    uploaded.save(saved_path)

    extracted_text, confidence = run_ocr_text(saved_path)

    scan = ScanHistory(filename=saved_name, recognized_text=extracted_text, confidence=confidence)
    db.session.add(scan)
    db.session.commit()

    return redirect(url_for("results", scan_id=scan.id))

@app.route("/export-docx/<int:scan_id>")
def export_docx(scan_id):
    scan = db.session.get(ScanHistory, scan_id)
    if not scan:
        return "Scan not found", 404
    
    doc = Document()
    doc.add_heading('Pashto OCR Extraction Results', 0)
    
    # Add text line by line to maintain formatting
    for line in scan.recognized_text.split('\n'):
        p = doc.add_paragraph(line)
        # Pashto is RTL
        p.alignment = 2 # Right align

    target = BytesIO()
    doc.save(target)
    target.seek(0)
    
    return send_file(
        target,
        as_attachment=True,
        download_name=f"Pashto_OCR_{scan_id}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route("/history")
def history():
    scans_db = ScanHistory.query.order_by(ScanHistory.timestamp.desc()).all()
    scans = [scan_to_view(scan) for scan in scans_db]
    return render_template("history.html", scans=scans)

@app.route("/results")
@app.route("/results/<int:scan_id>")
def results(scan_id: int | None = None):
    selected_db = None
    if scan_id is not None:
        selected_db = db.session.get(ScanHistory, scan_id)
    if selected_db is None:
        selected_db = ScanHistory.query.order_by(ScanHistory.timestamp.desc()).first()

    if selected_db is None:
        return redirect(url_for("new_scan"))

    selected = scan_to_view(selected_db)
    return render_template("results.html", scan=selected, text_result=selected["text"])

@app.route("/save-edits", methods=["POST"])
def save_edits():
    data = request.json
    scan_id = data.get("scan_id")
    edited_text = data.get("text")
    
    if not scan_id or edited_text is None:
        return {"status": "error", "message": "Missing data"}, 400
        
    scan = db.session.get(ScanHistory, scan_id)
    if not scan:
        return {"status": "error", "message": "Scan not found"}, 404
        
    # Update DB
    scan.recognized_text = edited_text
    db.session.commit()
    
    # Generate Training Data
    try:
        global pipeline
        if pipeline is None:
            pipeline = FullPagePashtoRecognition()
            
        # Re-segment to get crops
        img_path = UPLOAD_DIR / scan.filename
        crops = pipeline.segmenter.segment_lines(str(img_path))
        
        # Clean lines from edited text
        lines = [l.strip() for l in edited_text.split('\n') if l.strip()]
        
        # We only save if the number of lines matches (to be safe for ground truth)
        if len(crops) == len(lines):
            user_data_dir = Path("USER_TRAINING_DATA")
            images_dir = user_data_dir / "images"
            images_dir.mkdir(parents=True, exist_ok=True)
            
            labels_file = user_data_dir / "labels.txt"
            
            # Check if we already saved this scan to avoid duplicates in labels.txt
            # For simplicity, we just append and rely on a 'clean' script later if needed,
            # but we can use a set of filenames to avoid easy duplicates.
            existing_files = set()
            if labels_file.exists():
                with open(labels_file, "r", encoding="utf-8") as f:
                    for line in f:
                        parts = line.split('\t')
                        if parts:
                            existing_files.add(parts[0])

            count = 0
            with open(labels_file, "a", encoding="utf-8") as f:
                for i, (crop, label) in enumerate(zip(crops, lines)):
                    crop_name = f"images/{scan_id}_line_{i}.jpg"
                    if crop_name not in existing_files:
                        cv2.imwrite(str(user_data_dir / crop_name), crop)
                        f.write(f"{crop_name}\t{label}\n")
                        count += 1
            
            if count > 0:
                return {"status": "success", "message": f"Saved {count} new lines for training!"}
            else:
                return {"status": "success", "message": "Text updated. Lines already in training set."}
        else:
            return {
                "status": "partial_success", 
                "message": f"Text updated in history, but couldn't sync for training. Lines mismatch: AI found {len(crops)} lines, but you provided {len(lines)}. Please ensure each line matches the image lines."
            }
            
    except Exception as e:
        print(f"Error saving training data: {e}")
        return {"status": "error", "message": f"DB Updated, but training export failed: {str(e)}"}, 500

@app.route("/train-model", methods=["POST"])
def train_model():
    try:
        # We run this synchronously for now to ensure user knows when it's done,
        # but in a real app this would be a background task.
        # Since it's local and short (10-15 epochs on small data), it should be okay for a minute.
        train_on_user_data(epochs=15)
        
        # Reload pipeline to use the new model
        global pipeline
        pipeline = FullPagePashtoRecognition()
        
        return {"status": "success", "message": "Model updated successfully!"}
    except Exception as e:
        print(f"Training Error: {e}")
        return {"status": "error", "message": str(e)}, 500

if __name__ == "__main__":
    initialize_storage_and_db()
    app.run(debug=True)
